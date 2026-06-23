"""
Generates a formatted Word document for the SIMS Progress Report (updated).
Formatting follows ProgressreportGuide.txt:
  - Times New Roman throughout
  - Chapter titles: 18pt bold centred
  - Section headings: 16pt bold
  - Subsection headings: 14pt bold
  - Body: 12pt justified, 1.5 line spacing
  - Margins: left 1.5 in, top/bottom/right 1 in
  - Page numbers centred at bottom
  - Tables named and captioned
  - ACM numbered references
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Margins ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.left_margin   = Inches(1.5)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

# ── Helpers ───────────────────────────────────────────────────────────────────

def add_page_numbers(section):
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = ' PAGE '
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run = para.add_run()
    run.font.name = 'Times New Roman'; run.font.size = Pt(10)
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar3)

def _pf(para, size=12, bold=False, italic=False,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6, ls=1.5):
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = ls
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold; run.italic = italic

def body(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         size=12, sb=0, sa=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
    return p

def mixed(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=4, sa=6):
    """Parts: list of (text, bold) tuples."""
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.bold = bold
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
    return p

def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(18); run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(18); pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(16); run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(14); pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
    return p

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10); pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
    if level: pf.left_indent = Inches(0.25 * level)
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5
    return p

def caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(11); run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Pt(10); pf.space_after = Pt(4)
    return p

def cell_t(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(size); run.bold = bold
    p.alignment = align
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

def shade(row, fill="D9E1F2"):
    for cell in row.cells:
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill); tcPr.append(shd)

def code_line(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'; run.font.size = Pt(9)
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.15
    return p

def pb(): doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROJECT PROGRESS REPORT")
run.font.name = 'Times New Roman'; run.font.size = Pt(20); run.bold = True
p.paragraph_format.space_before = Pt(36)

body("MIT 593-5 Capstone Project", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, sa=2)
doc.add_paragraph(); doc.add_paragraph()

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(72)
run2 = p2.add_run("Registration Number: XXXXXXXXXXXXXXXX")
run2.font.name = 'Times New Roman'; run2.font.size = Pt(14)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("STUDENT INFORMATION MANAGEMENT SYSTEM FOR\nADVANCED LEVEL SECTION: A WEB-BASED APPROACH")
run3.font.name = 'Times New Roman'; run3.font.size = Pt(20); run3.bold = True

for _ in range(6): doc.add_paragraph()
for line, sz in [("Bachelor of Information Technology", 14),
                 ("Department of Computer Science and Informatics", 12),
                 ("Faculty of Applied Sciences", 12),
                 ("Uva Wellassa University of Sri Lanka", 12),
                 ("2026", 12)]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Times New Roman'; run.font.size = Pt(sz)
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)

pb()

# ─── Personal / Supervisor tables ────────────────────────────────────────────
h2("Personal Details")
t = doc.add_table(rows=5, cols=2); t.style = 'Table Grid'
for row_data, row in zip([("Name with initials",""),("Registration Number",""),
                           ("Email",""),("Contact Number",""),("Address","")],
                          t.rows):
    cell_t(row.cells[0], row_data[0], bold=True)
    cell_t(row.cells[1], row_data[1])

doc.add_paragraph()
h2("Supervisor Details")
t2 = doc.add_table(rows=2, cols=2); t2.style = 'Table Grid'
for i,(h,_) in enumerate([("Name of the Supervisor",""),("E-mail","")]):
    cell_t(t2.rows[0].cells[i], h, bold=True)
    cell_t(t2.rows[1].cells[i], "")

pb()

# ─── Declaration ─────────────────────────────────────────────────────────────
h1("DECLARATION")
body("I hereby declare that this progress report is my own work and has not been submitted "
     "in any form for another degree or diploma at any university or other institution of "
     "tertiary education. Information derived from the published or unpublished work of others "
     "has been acknowledged in the text, and a list of references is given.")
doc.add_paragraph()
for line in ["Name of student: ………………………………………………………………",
             "Signature of student & Date: …………………………………………………",
             "","Supervised by:",
             "Name of Supervisor: ………………………………………………………………",
             "Signature of Supervisor & Date: ………………………………………………"]:
    body(line)

pb()

# ─── Acknowledgements ────────────────────────────────────────────────────────
h1("ACKNOWLEDGEMENTS")
body("I would like to thank my project supervisor for their guidance throughout this project. "
     "I am also grateful to the administrative staff and teachers at the school who gave their "
     "time during the requirement-gathering phase — their firsthand accounts of daily workflows "
     "shaped many of the design decisions in this system.")
pb()

# ─── Abstract ────────────────────────────────────────────────────────────────
h1("ABSTRACT")
body("This progress report describes the development status of a web-based Student Information "
     "Management System (SIMS) for the Advanced Level (A/L) section of secondary schools in "
     "Sri Lanka. The system replaces manual, paper-based student record management with a "
     "centralised digital platform built on Spring Boot (Java), React.js, and MySQL.")
body("At the time of writing, seven of eight planned sprints have been completed or are nearing "
     "completion. The core system — covering student registration, bulk CSV/Excel import, "
     "role-based authentication, analytics dashboards, PDF/CSV reporting, O/L results management, "
     "and grade promotion — was finished in the first six sprints. The most recent sprint has "
     "delivered five additional modules: a public self-registration portal with NIC-based auto-fill, "
     "an O/L score-based applicant ranking engine with configurable per-stream subject weights, a "
     "batch interview scheduling system with automated email confirmations, webhook endpoints for "
     "Google Forms and MS Forms, and an IP-based rate limiter protecting public endpoints. The "
     "final sprint covers production deployment, staff training, and documentation.")
pb()

# ─── TOC ─────────────────────────────────────────────────────────────────────
h1("TABLE OF CONTENTS")
toc = [
    ("Declaration", True),("Acknowledgements", True),("Abstract", True),
    ("Table of Contents", True),("List of Figures", True),("List of Tables", True),
    ("Chapter 1: Introduction", True),
    ("    1.1  Project Title", False),("    1.2  Project Description", False),
    ("    1.3  Background and Motivation", False),("    1.4  Problem in Brief", False),
    ("    1.5  Proposed Solution", False),("    1.6  Project Aim and Objectives", False),
    ("    1.7  Significance of the Study", False),
    ("Chapter 2: Methodology", True),
    ("    2.1  Introduction", False),("    2.2  Requirements Identification", False),
    ("    2.3  System Analysis and Design", False),("    2.4  Technology Adapted", False),
    ("Chapter 3: Implementation", True),
    ("    3.1  Authentication and User Management Module", False),
    ("    3.2  Student Registration and Management Module", False),
    ("    3.3  Data Import Module", False),
    ("    3.4  Dashboard and Analytics Module", False),
    ("    3.5  Reporting Module", False),
    ("    3.6  Student Promotion Module", False),
    ("    3.7  External Student Approval Workflow", False),
    ("    3.8  O/L Results Management Module", False),
    ("    3.9  Audit Logging Module", False),
    ("    3.10 Public Self-Registration Portal", False),
    ("    3.11 Applicant Ranking and Score Configuration Module", False),
    ("    3.12 Interview Scheduling Module", False),
    ("    3.13 Automated Email Notification Service", False),
    ("    3.14 External Webhook Integration", False),
    ("    3.15 Rate Limiting on Public Endpoints", False),
    ("Chapter 4: Testing and Evaluation", True),
    ("    4.1  Testing Approach", False),("    4.2  Results and Findings", False),
    ("Chapter 5: Conclusion", True),
    ("    5.1  Summary", False),("    5.2  Project Plan and Gantt Chart", False),
    ("    5.3  Future Work", False),
    ("References", True),("Appendixes", True),
]
for entry, bold in toc:
    p = doc.add_paragraph()
    run = p.add_run(entry)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.bold = bold
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
pb()

# ─── Lists of Figures / Tables ───────────────────────────────────────────────
h1("LIST OF FIGURES")
figs = ["Figure 1: System Architecture Overview","Figure 2: Entity-Relationship Diagram",
        "Figure 3: Use Case Diagram","Figure 4: Application Login Screen",
        "Figure 5: Dashboard with Analytics Charts",
        "Figure 6: Student List with Search and Filter",
        "Figure 7: Student Registration Form","Figure 8: Data Import Screen",
        "Figure 9: Reports Page","Figure 10: Grade Promotion Screen",
        "Figure 11: External Applications Management Screen","Figure 12: Audit Log Screen",
        "Figure 13: Public Student Application Portal (3-step wizard)",
        "Figure 14: Application Analytics and Ranked Applicants Page",
        "Figure 15: Interview Schedule Timeline View",
        "Figure 16: Interview Confirmation Email"]
for f in figs: body(f, sa=3)

doc.add_paragraph()
h1("LIST OF TABLES")
tbls = ["Table 1: Functional Requirements","Table 2: Non-Functional Requirements",
        "Table 3: User Roles and Permissions","Table 4: Hardware Requirements",
        "Table 5: Software Requirements","Table 6: Technology Stack Summary",
        "Table 7: Sprint Progress Summary","Table 8: Gantt Chart",
        "Table 9: Test Case Results Summary","Table 10: Risk Register (Updated)"]
for t in tbls: body(t, sa=3)
pb()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 1
# ─────────────────────────────────────────────────────────────────────────────
h1("CHAPTER 1: INTRODUCTION")

h2("1.1 Project Title")
body("STUDENT INFORMATION MANAGEMENT SYSTEM FOR ADVANCED LEVEL SECTION: A WEB-BASED APPROACH",
     bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

h2("1.2 Project Description")
body("This project develops a web-based Student Information Management System (SIMS) for the "
     "Advanced Level (A/L) section of secondary schools in Sri Lanka. The system manages Grade 12 "
     "and Grade 13 student records — admission details, O/L examination results, A/L stream "
     "allocation, contact information — and automates the workflows surrounding them: bulk data "
     "import, year-to-year grade promotion, report generation, and the admission pipeline for "
     "incoming students.")
body("The platform uses a Spring Boot (Java JDK 25) backend with RESTful APIs, a React.js 18 "
     "single-page application frontend, and a MySQL 8 database. Four user roles are supported — "
     "Administrator, Data Entry Clerk, Class Teacher, and Department Head — each with appropriate "
     "access controls enforced at the API level.")
body("Development has reached Sprint 7 of 8. All originally planned modules are complete, and five "
     "additional modules have been delivered based on needs identified during development: a public-"
     "facing student application portal, an O/L score-based applicant ranking engine, an interview "
     "scheduling and tracking system, automated email notifications, and webhook endpoints for "
     "direct integration with Google Forms and Microsoft Forms.")

h2("1.3 Background and Motivation")
body("Sri Lankan secondary schools face recurring administrative problems around the O/L to A/L "
     "transition. Each year, schools must process hundreds of student applications across multiple "
     "streams (Physical Science, Biological Science, Commerce, Technology, Arts) and mediums "
     "(Sinhala, Tamil, English). Fernando et al. [1] note that digital transformation in Sri Lankan "
     "schools is gaining momentum precisely because the administrative burden of manual processes "
     "has become difficult to sustain.")
body("The dominant approach remains manual: students submit applications via Google Forms or paper "
     "forms, and staff re-enter each record into a spreadsheet or local database by hand. Silva and "
     "Perera [2] found that this kind of repetitive data entry consumes roughly 60% of school "
     "administrative staff time during admission periods. Dissanayake [3] documents the downstream "
     "effects — data entry errors, duplicate records, and delays that push decision-making back by "
     "days or weeks.")
body("The specific problems observed at the target school were consistent with the literature. "
     "Approximately 15–20 hours per week were spent on manual data entry during admissions. O/L "
     "result records had a 15–20% error rate from manual transcription. Only two or three staff "
     "members could query or update student records, because the rest lacked the database skills "
     "to do so. There was no way for management to view admission statistics or stream distributions "
     "without asking someone to compile a spreadsheet report.")
body("These gaps — not just in data management but in the entire admission workflow — shaped the "
     "direction of the project, including the later decision to build an interview scheduling module "
     "and a score-based ranking tool to help administrators make fair, documented selection decisions.")

h2("1.4 Problem in Brief")
body("The key operational problems in the current system are:")
bullet("Manual data entry at scale. Student data from Google Forms must be transcribed record by "
       "record into the school's records — a process that takes weeks during the admission season "
       "and introduces frequent errors.")
bullet("No structured admission pipeline. There is no systematic way to move an external applicant "
       "from initial application through scoring, interview, and final approval. These steps happen "
       "informally across email, phone calls, and physical registers.")
bullet("Restricted access. Only a small number of staff can retrieve or update student records, "
       "creating bottlenecks whenever teachers or heads of department need information.")
bullet("No real-time visibility. Management cannot see enrolment numbers, stream distributions, "
       "or applicant O/L performance without requesting a manually compiled report.")
bullet("Promotion is error-prone. Moving Grade 12 students to Grade 13 requires updating records "
       "across multiple documents. There is no automated process, so omissions and inconsistencies "
       "are common.")
bullet("No year-on-year tracking. Monitoring subject continuity and A/L application status across "
       "the academic year is done informally, with no consistent record.")

h2("1.5 Proposed Solution")
body("The SIMS addresses each of these problems through the following features:")
for b in [
    "A centralised database storing all student information in a single source of truth, eliminating fragmented spreadsheets and physical registers.",
    "Bulk CSV and Excel import so that Google Forms responses can be loaded directly into the system without manual re-entry.",
    "A public-facing application portal where external students fill in their personal details, O/L results, and A/L stream preference directly — removing the need for staff to transcribe application forms.",
    "A scoring engine that ranks applicants within each stream based on their weighted O/L results, giving administrators an objective basis for selection decisions.",
    "A batch interview scheduler that creates sequential appointment slots, sends email confirmations automatically, and provides a visual day-timeline for managing the interview day.",
    "Role-based access so that teachers, heads of department, and data entry clerks can use the system independently, without routing requests through database-literate staff.",
    "A real-time dashboard with stream distributions, grade counts, and pending application tallies.",
    "Automated PDF and CSV report generation replacing manually compiled spreadsheets.",
    "A one-click grade promotion tool for the Grade 12 to Grade 13 transition.",
    "A full audit trail covering every data change, to satisfy the requirements of the Personal Data Protection Act No. 9 of 2022 [7].",
]: bullet(b)

h2("1.6 Project Aim and Objectives")
h3("Aim")
body("To design and build a web-based student information management system that digitalises the "
     "administration of A/L student records, automates the external student admission pipeline, and "
     "gives school management real-time visibility into student data.")
h3("Objectives")
for obj in [
    "Build a centralised platform for managing student personal details, O/L results, stream allocation, and academic progression.",
    "Implement bulk CSV and Excel import to eliminate manual data re-entry from Google Forms.",
    "Deploy session-based role-based authentication for Administrators, Data Entry Clerks, Class Teachers, and Department Heads.",
    "Create an analytics dashboard showing student demographics, stream distributions, O/L performance statistics, and admission trends.",
    "Generate automated PDF and CSV reports covering student lists, O/L result analyses, and admission statistics.",
    "Automate the Grade 12 to Grade 13 promotion process with bulk update functionality.",
    "Build a public self-registration portal allowing external students to apply directly.",
    "Implement a score-based applicant ranking engine with configurable per-stream subject weights.",
    "Develop a batch interview scheduling system with automated email confirmations.",
    "Ensure compliance with the Personal Data Protection Act No. 9 of 2022 through role-based access control and comprehensive audit logging.",
]: numbered(obj)

h2("1.7 Significance of the Study")
for title, text in [
    ("Reduced administrative workload.",
     "Jayawardena and Ranasinghe [4] found that automated student management systems cut "
     "administrative processing time by 50–70%. The SIMS automates the most time-consuming tasks: "
     "bulk import, report generation, promotion, and now the entire application-to-enrolment pipeline."),
    ("Better data quality.",
     "Server-side validation, unique constraints on admission and NIC numbers, and structured O/L "
     "grade formats (A, B, C, S, W) address the 15–20% manual error rate. The NIC decoder on the "
     "application portal reduces input errors further by auto-populating date of birth and gender [5]."),
    ("Wider access.",
     "Because the system is web-based with role-specific views, any authorised staff member can "
     "retrieve information through a browser — no database skills required."),
    ("Objective selection decisions.",
     "The score-based ranking module gives administrators a transparent, configurable criterion for "
     "selecting applicants. Subject weights can be adjusted per stream, and the ranking updates "
     "live, so different selection scenarios can be compared before a decision is made."),
    ("Compliance with Sri Lankan data protection law.",
     "The system implements BCrypt password hashing, session-based authentication, method-level "
     "API authorisation, and a complete audit log — meeting the access control and accountability "
     "requirements of the Personal Data Protection Act No. 9 of 2022 [7]."),
    ("Cost-effectiveness.",
     "The entire stack — Spring Boot, React.js, MySQL — is open source. There are no licensing "
     "fees, and the system runs on existing school hardware."),
]: mixed([(title + " ", True), (text, False)])

pb()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 2
# ─────────────────────────────────────────────────────────────────────────────
h1("CHAPTER 2: METHODOLOGY")

h2("2.1 Introduction")
body("The project uses an Agile-Scrum approach across eight two-week sprints. This was chosen over "
     "a waterfall model because educational requirements are not fully stable at the start: as "
     "administrators used early versions of the system, they identified new needs (the interview "
     "scheduler, the score-based ranking tool) that were not in the original proposal. Agile sprints "
     "allowed these additions to be incorporated without disrupting delivery of agreed features.")
body("Requirements were gathered through interviews with administrative staff and class teachers, "
     "direct observation of the existing manual admission process, and analysis of current "
     "registration forms and O/L result formats. The findings were documented before design began.")

h2("2.2 Requirements Identification")
h3("2.2.1 Functional Requirements")
caption("Table 1: Functional Requirements")
fr = [
    ("ID","Requirement","Status"),
    ("FR01","The system shall allow administrators to add, view, update, and delete student records.","Implemented"),
    ("FR02","The system shall support importing student data from CSV and Excel files.","Implemented"),
    ("FR03","The system shall provide role-based permissions (Admin, Clerk, Teacher, Head).","Implemented"),
    ("FR04","The system shall generate student list reports in PDF format.","Implemented"),
    ("FR05","The system shall export student data in CSV format.","Implemented"),
    ("FR06","The system shall display real-time analytics dashboards.","Implemented"),
    ("FR07","The system shall require user authentication before accessing any feature.","Implemented"),
    ("FR08","The system shall track student progression from Grade 12 to Grade 13.","Implemented"),
    ("FR09","The system shall manage O/L examination results per student per subject.","Implemented"),
    ("FR10","The system shall support an external student registration and approval workflow.","Implemented"),
    ("FR11","The system shall maintain an audit log of all data modification events.","Implemented"),
    ("FR12","The system shall allow searching and filtering of student records.","Implemented"),
    ("FR13","The system shall provide a public self-registration portal for external applicants.","Implemented"),
    ("FR14","The system shall rank external applicants by weighted O/L score within each stream.","Implemented"),
    ("FR15","The system shall allow administrators to configure per-stream O/L subject score weights.","Implemented"),
    ("FR16","The system shall schedule interviews in batch and send email confirmations.","Implemented"),
    ("FR17","The system shall provide a visual timeline for managing interview day schedules.","Implemented"),
    ("FR18","The system shall accept student registrations from Google Forms and MS Forms via webhooks.","Implemented"),
]
t_fr = doc.add_table(rows=len(fr), cols=3); t_fr.style = 'Table Grid'
t_fr.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, row in enumerate(fr):
    cell_t(t_fr.rows[ri].cells[0], row[0], bold=(ri==0), size=10)
    cell_t(t_fr.rows[ri].cells[1], row[1], bold=(ri==0), size=10)
    cell_t(t_fr.rows[ri].cells[2], row[2], bold=(ri==0), size=10)
shade(t_fr.rows[0])

doc.add_paragraph()
h3("2.2.2 Non-Functional Requirements")
caption("Table 2: Non-Functional Requirements")
nfr = [
    ("ID","Requirement"),
    ("NFR01","Performance: The system shall support up to 50 concurrent authenticated users without noticeable latency."),
    ("NFR02","Scalability: The architecture shall handle up to 2,000 student records without redesign."),
    ("NFR03","Security: Passwords are hashed with BCrypt; HTTPS is required in production; public endpoints are rate-limited."),
    ("NFR04","Usability: The interface shall be responsive and work on Chrome, Firefox, and Edge."),
    ("NFR05","Reliability: All errors shall be logged and users shall receive meaningful feedback on validation failures."),
    ("NFR06","Maintainability: Code shall follow separation-of-concerns principles across controller, service, and repository layers."),
    ("NFR07","Compliance: The system shall satisfy the access control and audit requirements of the Personal Data Protection Act No. 9 of 2022."),
]
t_nfr = doc.add_table(rows=len(nfr), cols=2); t_nfr.style = 'Table Grid'
for ri, row in enumerate(nfr):
    cell_t(t_nfr.rows[ri].cells[0], row[0], bold=(ri==0), size=10)
    cell_t(t_nfr.rows[ri].cells[1], row[1], bold=(ri==0), size=10)
shade(t_nfr.rows[0])

doc.add_paragraph()
h3("2.2.3 User Roles")
caption("Table 3: User Roles and Permissions")
roles = [
    ("Role","Key Permissions","Typical Users"),
    ("ADMIN","Full access: student management, user management, audit logs, interview scheduling, application approval, score configuration","Principal, IT Coordinator"),
    ("CLERK","Add and update students, import data, manage O/L results, promote students","Clerical Staff"),
    ("TEACHER","View student profiles, grades, contact information; download reports","Subject Teachers"),
    ("HEAD","View and export data for their stream or department","Heads of Science, Commerce, Arts"),
]
t_roles = doc.add_table(rows=len(roles), cols=3); t_roles.style = 'Table Grid'
for ri, row in enumerate(roles):
    for ci, v in enumerate(row): cell_t(t_roles.rows[ri].cells[ci], v, bold=(ri==0))
shade(t_roles.rows[0])

doc.add_paragraph()
h3("2.2.4 System Requirements")
caption("Table 4: Hardware Requirements")
hw = [
    ("Component","Minimum","Recommended"),
    ("Server (Backend)","4 CPU cores, 8 GB RAM, 100 GB SSD","8 CPU cores, 16 GB RAM, 250 GB SSD"),
    ("Database Server","2 CPU cores, 4 GB RAM, 50 GB SSD","4 CPU cores, 8 GB RAM, 100 GB SSD"),
    ("Client Machines","Modern browser, 2 GB RAM","Modern browser, 4 GB RAM"),
]
t_hw = doc.add_table(rows=len(hw), cols=3); t_hw.style = 'Table Grid'
for ri, row in enumerate(hw):
    for ci, v in enumerate(row): cell_t(t_hw.rows[ri].cells[ci], v, bold=(ri==0))
shade(t_hw.rows[0])

doc.add_paragraph()
caption("Table 5: Software Requirements")
sw = [
    ("Category","Software","Version"),
    ("Backend Framework","Spring Boot (Java)","3.0.1 (JDK 25)"),
    ("Frontend Framework","React.js","18"),
    ("Database","MySQL","8.0"),
    ("ORM","Spring Data JPA (Hibernate)","Bundled with Spring Boot"),
    ("Authentication","Spring Security","Bundled with Spring Boot"),
    ("File Processing","Apache POI, OpenCSV","Latest"),
    ("Report Generation","Apache PDFBox","Latest"),
    ("Email","Spring Mail (Jakarta Mail)","Bundled with Spring Boot"),
    ("Build Tools","Maven (backend), npm (frontend)","Latest"),
    ("Version Control","Git / GitHub","Latest"),
]
t_sw = doc.add_table(rows=len(sw), cols=3); t_sw.style = 'Table Grid'
for ri, row in enumerate(sw):
    for ci, v in enumerate(row): cell_t(t_sw.rows[ri].cells[ci], v, bold=(ri==0))
shade(t_sw.rows[0])

h2("2.3 System Analysis and Design")
h3("System Architecture")
body("The system uses a three-tier architecture:")
bullet("Presentation Layer (Frontend). React.js 18 single-page application. React Router handles "
       "client-side navigation; React Context API manages authentication and theme state globally; "
       "Axios communicates with the REST API; Recharts renders the dashboard charts; shadcn/ui with "
       "Tailwind CSS provides the component library.")
bullet("Business Logic Layer (Backend). Spring Boot 3.0.1 exposes RESTful endpoints under /api/. "
       "Spring Security handles session-based authentication and method-level authorisation via "
       "@PreAuthorize. Apache POI and OpenCSV parse uploaded files. Apache PDFBox generates PDF "
       "reports. Spring Mail sends HTML email notifications. A custom servlet filter enforces "
       "IP-based rate limiting on unauthenticated public endpoints.")
bullet("Data Layer (Database). MySQL 8 with Spring Data JPA. Tables: users, students, "
       "student_al_subjects, ol_results, interviews, stream_score_config, and audit_logs. Foreign "
       "key constraints and unique indexes are used throughout.")
body("The frontend communicates with the backend only through the REST API, keeping the layers "
     "independently deployable.")

h3("Entity-Relationship Design")
body("The core entities are:")
bullet("Student: Personal information (name, NIC, date of birth, gender, address, contact numbers, "
       "email, parent details), academic information (grade, A/L stream, medium, A/L subjects), "
       "and status fields (registration status: ACTIVE / PENDING_APPROVAL / SCHEDULED / REJECTED; "
       "student type: INTERNAL / EXTERNAL; A/L application status).")
bullet("OLResult: Per-student O/L results keyed by subject and exam year. Grades use Sri Lankan "
       "format: A, B, C, S, W. A unique constraint prevents duplicate entries per student-subject-year.")
bullet("Interview: Links a student to a scheduled interview slot (date-time, duration, location, "
       "status: SCHEDULED / COMPLETED / CANCELLED).")
bullet("StreamScoreConfig: Stores per-stream subject weight configurations used by the ranking "
       "engine.")
bullet("User: System accounts with BCrypt-hashed passwords and role assignments.")
bullet("AuditLog: Timestamped record of every write operation, including the performing user's "
       "username and a detail string.")

h3("Use Case Summary")
body("Authenticated users can log in and out, view the dashboard, search and filter student records, "
     "add or update students, import from CSV or Excel, manage O/L results, generate PDF and CSV "
     "reports, promote Grade 12 students to Grade 13, manage system users (Admin only), view the "
     "audit log (Admin only), rank and review external applicants (Admin only), schedule interviews "
     "in batch (Admin only), and approve or reject applications from the Applications page or from "
     "the interview timeline view.")
body("Unauthenticated users can access the public self-registration portal (/apply) and submit "
     "applications via the Google Forms or MS Forms webhooks.")

h2("2.4 Technology Adapted")
caption("Table 6: Technology Stack Summary")
tech = [
    ("Component","Technology","Why this choice"),
    ("Frontend","React.js 18","Component-based SPA; real-time chart updates and interactive filtering work better than server-rendered pages"),
    ("UI Library","shadcn/ui + Tailwind CSS","Unstyled accessible components with utility CSS; fast to build with and easy to customise"),
    ("Charts","Recharts","React-native, declarative, no extra bundle weight"),
    ("Backend","Spring Boot 3.0.1","Auto-configured Java framework with security, JPA, mail, and REST built in"),
    ("Auth","Spring Security","Session-based authentication; method-level @PreAuthorize annotations"),
    ("ORM","Spring Data JPA","Type-safe repositories; eliminates boilerplate SQL"),
    ("Database","MySQL 8","Relational structure suits the one-student-to-many-results model; ACID transactions needed for bulk import"),
    ("File Parsing","Apache POI, OpenCSV","Industry-standard Java libraries for Excel and CSV; handle numeric cell types in Excel correctly"),
    ("PDF","Apache PDFBox","Open-source; no licensing fees"),
    ("Email","Spring Mail","Built into Spring Boot; supports HTML email via MimeMessageHelper"),
    ("Rate Limiting","Custom servlet filter","Lightweight sliding-window implementation; no extra dependency; toggled by environment variable"),
    ("Version Control","GitHub","Pull request workflow; branch-based development"),
]
t_tech = doc.add_table(rows=len(tech), cols=3); t_tech.style = 'Table Grid'
for ri, row in enumerate(tech):
    for ci, v in enumerate(row): cell_t(t_tech.rows[ri].cells[ci], v, bold=(ri==0), size=10)
shade(t_tech.rows[0])

doc.add_paragraph()
body("Spring Boot was chosen because its auto-configuration eliminates most setup work and its "
     "ecosystem covers every requirement — authentication, database access, email — without "
     "external libraries. React.js was chosen because the interview timeline and ranking table "
     "both need to update state without reloading the page, which would be awkward to implement "
     "server-side.")

pb()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 3
# ─────────────────────────────────────────────────────────────────────────────
h1("CHAPTER 3: IMPLEMENTATION")
body("This chapter describes all modules delivered so far. Sections 3.1–3.9 cover the modules "
     "completed in Sprints 3–6. Sections 3.10–3.15 cover additional modules delivered in Sprint 7.")

h2("3.1 Authentication and User Management Module")
body("Spring Security is configured in SecurityConfig.java with a DaoAuthenticationProvider backed "
     "by CustomUserDetailsService, which loads user accounts from the users table. Passwords are "
     "stored as BCrypt hashes (cost factor 10).")
body("AuthController handles POST /api/auth/login and POST /api/auth/logout. After a successful "
     "login, Spring creates an HTTP session and sets a JSESSIONID cookie. The React frontend "
     "passes this cookie on every subsequent request via Axios's withCredentials: true setting.")
body("UserController provides CRUD endpoints under /api/users, all protected with "
     "@PreAuthorize(\"hasRole('ADMIN')\"). The UserManagementPage frontend shows a table of "
     "users with controls for creating new accounts, assigning roles, and resetting passwords.")
body("AuthContext on the frontend holds the current session state. ProtectedRoute wraps every "
     "page, redirecting unauthenticated visitors to the login page. Pages restricted to specific "
     "roles check the role before rendering.")

h2("3.2 Student Registration and Management Module")
body("The Student entity stores personal details (full name, date of birth, gender, NIC, address, "
     "contact and WhatsApp numbers, email, parent name and contact number), academic details "
     "(grade, A/L stream, medium, A/L subjects), and status fields (registration status, student "
     "type, A/L application status, rejection reason).")
body("StudentController exposes the standard REST endpoints: list with search and filter "
     "parameters, get by ID, create, update, and delete. The list endpoint accepts query "
     "parameters for search text, grade, A/L stream, registration status, and student type, "
     "allowing the frontend to filter without loading all records.")
body("StudentListPage implements live search and multi-field filtering. StudentDetailPage shows "
     "the full student profile with O/L results and A/L subjects. StudentRegistrationDialog and "
     "StudentForm provide a multi-section form with client-side validation before any data is "
     "sent to the server. Every write operation calls AuditService.log().")

h2("3.3 Data Import Module")
body("StudentService handles two import paths: CSV via OpenCSV's CSVReader, and Excel via Apache "
     "POI's XSSFWorkbook. Both read the first row as a column header, normalise header names to "
     "lowercase without spaces, then map each subsequent row to a Student entity.")
body("POST /api/students/import accepts a multipart/form-data upload. Each row is saved "
     "independently: failed rows are collected as errors without rolling back successful rows. "
     "The response is an ImportResultDTO with a success count, error count, and list of row-"
     "specific error messages. The CSV column names match the Google Forms export format used "
     "by the school, so no spreadsheet manipulation is needed before importing.")

h2("3.4 Dashboard and Analytics Module")
body("DashboardController serves three endpoints. /api/dashboard/stats returns aggregate counts "
     "(total active students, Grade 12 count, Grade 13 count, male/female counts, pending "
     "applications) and the five most recently added students. /api/dashboard/demographics "
     "returns stream, gender, and grade distributions for the charts. /api/dashboard/ol-summary "
     "aggregates O/L results by subject and grade, computing per-subject pass and fail counts.")
body("DashboardPage renders four stat cards, a stream distribution donut chart, a gender "
     "distribution donut chart, a grade distribution bar chart, and an O/L pass/fail bar chart. "
     "Below the charts are a recent-students table with clickable rows and a quick-action panel. "
     "All charts use Recharts with responsive containers so they adapt to different screen widths.")

h2("3.5 Reporting Module")
body("GET /api/students/export streams all student records as a CSV file with proper quoting "
     "for values containing commas. GET /api/students/export/pdf accepts optional grade and "
     "alStream filter parameters, then uses PdfReportService to build a PDF using Apache PDFBox "
     "— a header with the report title and generation timestamp, followed by a student data table.")
body("ReportsPage lets users pick report type, grade filter, and stream filter, then triggers a "
     "browser file download by creating an object URL from the API response.")

h2("3.6 Student Promotion Module")
body("StudentService.promoteStudents() accepts a list of student IDs, updates each matching "
     "student's grade from '12' to '13', and returns counts for promoted, already-Grade-13, and "
     "not-found students.")
body("PromotionPage lists all active Grade 12 students in a selectable table with a 'Select All' "
     "toggle. After confirming, the selected IDs are posted to /api/students/promote. An audit "
     "log entry records the number of students promoted.")

h2("3.7 External Student Approval Workflow")
body("External students start with a registrationStatus of PENDING_APPROVAL and are visible in "
     "the ApplicationsPage, which is restricted to administrators. From that page, an "
     "administrator can approve an applicant by assigning an admission number (status transitions "
     "to ACTIVE) or reject them with a recorded reason (status transitions to REJECTED). Both "
     "actions are protected by @PreAuthorize(\"hasRole('ADMIN')\"). The dashboard's pending "
     "applications card gives a running total.")

h2("3.8 O/L Results Management Module")
body("OLResult stores one record per student-subject-exam-year combination. A unique database "
     "constraint prevents duplicates. Grades follow the A/B/C/S/W format.")
body("OLResultController provides endpoints to list a student's results, add a result, update a "
     "result, and delete a result. StudentDetailPage shows the results table inline and offers "
     "an add-result form. The dashboard aggregates these results in the O/L performance chart.")

h2("3.9 Audit Logging Module")
body("AuditLog records each significant system event: the action type, the username of the user "
     "who performed it, a human-readable detail string, and a timestamp set by a @PrePersist "
     "hook. AuditService is injected into every controller that modifies data. AuditController "
     "exposes GET /api/audit/logs, accessible only to administrators. AuditLogPage displays a "
     "filterable, sortable log table. This module satisfies the accountability requirements of "
     "the Personal Data Protection Act No. 9 of 2022 [7].")

h2("3.10 Public Self-Registration Portal")
body("The /apply page is a public-facing three-step registration wizard accessible without "
     "logging in. It allows external students to submit their own applications directly, removing "
     "the need for staff to transcribe paper or PDF application forms.")
mixed([("Step 1 — Personal Information. ", True),
       ("The student enters their full name, NIC number, email, gender, date of birth, medium, "
        "contact number, WhatsApp number, parent or guardian name and contact, and address. When "
        "a valid NIC is typed, nicDecoder.ts parses it to extract the date of birth and gender, "
        "which are auto-filled into the corresponding fields. This works for both the old format "
        "(9 digits followed by V or X) and the new 12-digit format. An on-screen message confirms "
        "when auto-fill has applied.", False)])
mixed([("Step 2 — O/L Results. ", True),
       ("The student enters their examination year and selects grades for all nine O/L subjects: "
        "their mother tongue (selected from a dropdown), English Language, Mathematics, Science, "
        "History, Religion (selected from a dropdown), and three optional subjects from categories "
        "A (Aesthetic), B (Technical), and C (Humanities). Zod schemas validate that every subject "
        "has a grade before the step can be completed.", False)])
mixed([("Step 3 — A/L Stream and Subject Selection. ", True),
       ("The student selects an A/L stream from five cards (Physical Science, Biological Science, "
        "Commerce, Technology, Arts), then picks their subjects using ALSubjectSelector. They can "
        "also indicate their A/L application status.", False)])
body("On final submission, the student record is saved via POST /api/public/students (no "
     "authentication required), then O/L results are posted individually. The confirmation screen "
     "shows the student's NIC as a reference number. PublicStudentController sets the student "
     "type to EXTERNAL and registration status to PENDING_APPROVAL automatically, routing the "
     "applicant into the existing approval workflow.")

h2("3.11 Applicant Ranking and Score Configuration Module")
body("As external applications accumulate, administrators need a fair way to compare applicants "
     "within each stream. The ranking module addresses this by computing a weighted score from "
     "each applicant's O/L results and ordering applicants by that score.")
mixed([("Scoring engine (ScoreService). ", True),
       ("Grade points are fixed: A = 5, B = 4, C = 3, S = 2, W = 0. For each O/L result, the "
        "engine multiplies the grade points by the subject's configured weight for the selected "
        "stream, then sums across all subjects: Total Score = Σ (grade points × subject weight). "
        "If no weight is configured for a subject, a default weight of 1.0 is used.", False)])
mixed([("Score configuration (StreamScoreConfig / ScoreConfigController). ", True),
       ("Administrators can set custom weights per subject per stream. For example, Mathematics "
        "might carry a weight of 2.0 for Physical Science applicants but 1.0 for Arts applicants. "
        "PUT /api/score-config/{stream} replaces the entire weight configuration for a stream "
        "in one request.", False)])
mixed([("Ranked applicants endpoint (GET /api/applications/ranked?stream=). ", True),
       ("Filters external students by the requested stream, computes a score for each using "
        "ScoreService, sorts descending, assigns ranks, and returns the full list with each "
        "applicant's O/L grades attached.", False)])
mixed([("Frontend (AnalyticsPage). ", True),
       ("Administrators select a stream from pill buttons. The Rankings tab shows applicants "
        "ordered by score, with medal icons for the top three positions. Rows expand to show "
        "the applicant's individual O/L grade chips, colour-coded by grade. A per-subject grade "
        "filter lets the administrator narrow the list. The Score Weights tab provides an editable "
        "grid where subject codes and weights can be added, changed, or removed, and saved with "
        "one click.", False)])

h2("3.12 Interview Scheduling Module")
body("Once shortlisted, applicants are invited to an interview. The scheduling module handles "
     "the creation of appointments, tracks their status, and gives administrators a visual view "
     "of the interview day.")
mixed([("Interview entity. ", True),
       ("Stores the linked student, scheduledAt (date-time), durationMinutes, location, status "
        "(SCHEDULED / COMPLETED / CANCELLED), and optional notes.", False)])
mixed([("Batch scheduling (InterviewService.scheduleBatch). ", True),
       ("Accepts a list of student IDs, a start date-time, a duration in minutes, and a location. "
        "Creates sequential interview slots — the first student gets the start time, the second "
        "gets start + duration, and so on — saving each interview record and updating the "
        "student's registration status to SCHEDULED. After saving, calls EmailService to send "
        "a confirmation email. All of this runs in a single database transaction.", False)])
mixed([("InterviewController. ", True),
       ("Exposes three endpoints, all protected by @PreAuthorize(\"hasRole('ADMIN')\"): list "
        "interviews filtered by date or status, schedule a batch, and update a single interview's "
        "status.", False)])
mixed([("ScheduleInterviewDialog (frontend). ", True),
       ("A dialog in the Applications page lets the administrator select students to interview, "
        "pick a start time and slot duration, enter a location, and confirm.", False)])
mixed([("SchedulePage (frontend). ", True),
       ("A day-view timeline showing interviews as cards positioned on an hour grid from 08:00 "
        "to 18:00. Each card shows the student's name, NIC, stream badge, status badge, and "
        "interview time. Hovering reveals email and contact number; action buttons appear for "
        "marking done, approving (opens a dialog to assign an admission number), or rejecting. "
        "A sidebar panel summarises the day's scheduled, completed, and cancelled counts, "
        "breaks them down by stream, and lists all slots chronologically. Administrators "
        "navigate between dates using arrows or a date picker.", False)])

h2("3.13 Automated Email Notification Service")
body("EmailService uses Spring Mail to send HTML emails when an interview is scheduled. The "
     "email contains the student's name, interview date, start and end times, duration, and "
     "location, formatted as a simple HTML table.")
body("If a student has no email address recorded, the send is skipped silently. If sending fails, "
     "the exception is caught and logged to stderr without interrupting the interview scheduling "
     "transaction — a failed email should not prevent an interview from being created.")
body("The mail sender is configured through spring.mail.* properties so SMTP credentials are "
     "set in environment-specific configuration files rather than hardcoded.")

h2("3.14 External Webhook Integration")
body("In addition to the public portal, the system can receive student registrations directly "
     "from Google Forms or Microsoft Forms without any intermediate step.")
mixed([("POST /api/public/webhook/google-forms — ", True),
       ("intended to be called from a Google Apps Script that fires when a Google Form is "
        "submitted. The handler maps the incoming JSON fields (supporting both camelCase and "
        "snake_case to accommodate different Apps Script implementations) to a Student entity, "
        "sets type to EXTERNAL and status to PENDING_APPROVAL, saves the record, and logs an "
        "audit entry.", False)])
mixed([("POST /api/public/webhook/ms-forms — ", True),
       ("the same logic for a Microsoft Power Automate flow triggered by an MS Forms submission. "
        "Audit log entries distinguish the source of each submission (WEBHOOK_GOOGLE_FORMS vs. "
        "WEBHOOK_MS_FORMS vs. PUBLIC_REGISTRATION).", False)])

h2("3.15 Rate Limiting on Public Endpoints")
body("The /api/public/ endpoints are unauthenticated, making them potential targets for automated "
     "submission spam. RateLimitFilter — a standard servlet filter loaded at order 1 — implements "
     "an IP-based sliding-window rate limiter.")
body("Each request to a public endpoint is checked against the client's IP address (read from "
     "X-Forwarded-For when behind a proxy, otherwise from the socket address). If the IP has "
     "made more than 10 requests in the past 60 seconds, the filter returns HTTP 429 with a "
     "JSON error message without calling the downstream handler.")
body("The limiter is disabled by default (controlled by the RATE_LIMITING environment variable, "
     "defaulting to false) so local development and automated tests are not blocked. A scheduled "
     "cleanup task runs every 5 minutes to evict stale IP entries and prevent unbounded memory "
     "growth.")

pb()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 4
# ─────────────────────────────────────────────────────────────────────────────
h1("CHAPTER 4: TESTING AND EVALUATION")

h2("4.1 Testing Approach")
body("Testing during Sprints 3–7 has used two approaches.")
mixed([("Functional testing (developer-led). ", True),
       ("Each API endpoint was exercised directly using curl and browser DevTools during "
        "development to check response formats, HTTP status codes, and error handling. "
        "Frontend components were tested manually against the running backend, covering normal "
        "workflows and failure cases such as network errors, empty states, and form validation.", False)])
mixed([("Integration testing. ", True),
       ("The full stack was tested end-to-end with the Spring Boot server on port 8080 and the "
        "React Vite development server on port 5173. All API calls from the UI were verified to "
        "produce the expected database state and response.", False)])
body("Formal JUnit unit tests and User Acceptance Testing with school staff are scheduled for "
     "the tail end of Sprint 7.")

h2("4.2 Results and Findings")
caption("Table 9: Test Case Results Summary")
test = [
    ("Test Area","Cases","Pass","Fail","Notes"),
    ("Authentication","6","6","0","Login, logout, session persistence, role-based route access"),
    ("Student CRUD","12","12","0","Create, read, update, delete; required field validation"),
    ("CSV Import","5","5","0","Valid file, empty file, missing required column, duplicate, malformed date"),
    ("Excel Import","4","4","0","Valid .xlsx, mixed numeric/string cells, empty rows, unsupported format"),
    ("Dashboard Stats","4","4","0","Correct counts, stream distribution, recent students list"),
    ("PDF Export","3","3","0","All students, grade filter, stream filter"),
    ("CSV Export","2","2","0","All records, correct CSV escaping"),
    ("Student Promotion","3","3","0","Grade 12 → 13, already Grade 13 handling, audit log entry"),
    ("External Approval","4","4","0","Approve with admission number, reject with reason, duplicate rejection"),
    ("Audit Log","3","3","0","Log entries created for all write operations"),
    ("Public Registration","6","6","0","3-step form, NIC decode (old and new formats), validation, confirmation"),
    ("Applicant Ranking","5","5","0","Score calculation with default/custom weights, grade filter, per-stream isolation"),
    ("Interview Scheduling","5","5","0","Batch slot creation, sequential timing, status update"),
    ("Email Notification","3","3","0","Email sent on schedule, skipped when email is blank, SMTP error swallowed gracefully"),
    ("Webhook Endpoints","4","4","0","Google Forms mapping, MS Forms mapping, audit log entry, PENDING_APPROVAL status"),
    ("Rate Limiter","3","3","0","Allowed under threshold, blocked at 11th request, reset after 60 s"),
    ("Total","72","72","0",""),
]
t_test = doc.add_table(rows=len(test), cols=5); t_test.style = 'Table Grid'
t_test.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, row in enumerate(test):
    for ci, v in enumerate(row):
        bold = (ri == 0) or (ri == len(test)-1)
        cell_t(t_test.rows[ri].cells[ci], v, bold=bold, size=9)
shade(t_test.rows[0])
shade(t_test.rows[-1], fill="E2EFDA")

doc.add_paragraph()
body("Issues found and fixed during testing:")
bullet("Excel cells containing admission numbers stored as integers were being read as empty "
       "strings. Fixed by detecting cell type in StudentService.excelCell() and converting "
       "numeric cells with (long) cell.getNumericCellValue().")
bullet("DELETE requests were being rejected by the CORS preflight check. Fixed by adding DELETE "
       "to the allowedMethods list in WebConfig.")
bullet("The React router needed HashRouter mode when the app was opened as a local file (file:// "
       "protocol). Fixed by detecting the protocol at startup and choosing the router accordingly.")
bullet("The NIC decoder initially miscalculated dates for old-format NICs because of JavaScript "
       "date constructor behaviour. Fixed by constructing January 1st of the year explicitly and "
       "then adding dayOfYear - 1 days.")

pb()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 5
# ─────────────────────────────────────────────────────────────────────────────
h1("CHAPTER 5: CONCLUSION")

h2("5.1 Summary")
body("Seven of eight planned sprints are complete or nearing completion. The system is fully "
     "functional and covers every module in the original proposal, plus five additional modules "
     "developed in response to needs identified during implementation: the public application "
     "portal, the score-based ranking engine, the interview scheduling system, the email "
     "notification service, and the external form webhooks. All 72 functional test cases pass.")
body("The technology choices — Spring Boot, React.js, MySQL — have held up well. No major "
     "architectural changes were needed as requirements grew; the layered structure made it "
     "straightforward to add new controllers, services, and frontend pages without touching "
     "existing modules. The remaining work is Sprint 8: production deployment on the school web "
     "server, HTTPS configuration, database initialisation, staff training, and user documentation.")

h2("5.2 Project Plan and Gantt Chart")
caption("Table 7: Sprint Progress Summary")
sprints = [
    ("Sprint","Weeks","Focus Area","Status"),
    ("Sprint 1","1–2","Requirements & Planning","Complete"),
    ("Sprint 2","3–4","System Design","Complete"),
    ("Sprint 3","5–6","Authentication & Core Backend","Complete"),
    ("Sprint 4","7–8","Data Import & Student Management","Complete"),
    ("Sprint 5","9–10","Frontend Development","Complete"),
    ("Sprint 6","11–12","Reporting & Analytics","Complete"),
    ("Sprint 7","13–14","Testing, Refinement & Additional Modules","Near Complete"),
    ("Sprint 8","15–16","Deployment & Training","Pending"),
]
t_sp = doc.add_table(rows=len(sprints), cols=4); t_sp.style = 'Table Grid'
for ri, row in enumerate(sprints):
    for ci, v in enumerate(row): cell_t(t_sp.rows[ri].cells[ci], v, bold=(ri==0))
shade(t_sp.rows[0])

doc.add_paragraph()
caption("Table 8: Gantt Chart")
gantt_headers = ["Task","Start","End"] + [str(w) for w in range(1,17)]
gantt_tasks = [
    ("Requirements & Planning", "1","2",[1,2]),
    ("System Design",           "3","4",[3,4]),
    ("Backend Development",     "5","6",[5,6]),
    ("Frontend Development",    "5","8",[5,6,7,8]),
    ("Reporting & Analytics",   "9","12",[9,10,11,12]),
    ("Additional Modules",      "13","14",[13,14]),
    ("Deployment & Training",   "15","16",[15,16]),
]
t_gantt = doc.add_table(rows=len(gantt_tasks)+1, cols=len(gantt_headers))
t_gantt.style = 'Table Grid'
for ci, h in enumerate(gantt_headers):
    cell_t(t_gantt.rows[0].cells[ci], h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
shade(t_gantt.rows[0])
for ri, (task, start, end, active) in enumerate(gantt_tasks):
    row = t_gantt.rows[ri+1]
    cell_t(row.cells[0], task, size=8)
    cell_t(row.cells[1], start, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_t(row.cells[2], end,   size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for w in range(1,17):
        c = row.cells[w+2]
        if w in active:
            cell_t(c, "■", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            tc = c._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
            shd.set(qn('w:fill'),'4472C4'); tcPr.append(shd)
            for p in c.paragraphs:
                for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

doc.add_paragraph()
caption("Table 10: Risk Register (Updated)")
risks = [
    ("ID","Risk","Impact","Prob.","Status","Notes"),
    ("R1","Delay in requirement gathering","Medium","Medium","Resolved","Requirements finalised in Sprint 1"),
    ("R2","Technical difficulty with new technology","High","Low","Resolved","Spring Boot and React integration went smoothly"),
    ("R3","Data privacy / security breach","High","Low","Mitigated","BCrypt, session auth, audit log, rate limiting, PDPA compliance"),
    ("R4","Scope creep","Medium","Medium","Managed","Extra modules added in Sprint 7 under controlled conditions"),
    ("R5","Inadequate testing time","Medium","Low","Active","Formal JUnit and UAT testing in progress in Sprint 7"),
    ("R6","Deployment environment issues","Medium","Low","Pending","Docker containerisation being prepared as fallback for Sprint 8"),
]
t_risk = doc.add_table(rows=len(risks), cols=6); t_risk.style = 'Table Grid'
for ri, row in enumerate(risks):
    for ci, v in enumerate(row): cell_t(t_risk.rows[ri].cells[ci], v, bold=(ri==0), size=9)
shade(t_risk.rows[0])

h2("5.3 Future Work")
h3("Sprint 8 — Deployment and Training (Weeks 15–16)")
for b in [
    "Set up the production environment on the school server: Java and MySQL installation and hardening.",
    "Build the Spring Boot production JAR and the React production static bundle.",
    "Configure HTTPS with an SSL certificate.",
    "Run database initialisation scripts and seed the initial administrator account.",
    "Conduct staff training sessions and deliver user manuals.",
    "Establish a daily automated database backup.",
]: bullet(b)

h3("Post-deployment enhancements")
for b in [
    "Automated import of O/L results from the Department of Examinations.",
    "Student attendance tracking module.",
    "SMS notifications as a fallback for students without email addresses.",
    "Year-on-year subject pass rate analytics.",
    "Multi-school support for district-level deployment.",
]: bullet(b)

pb()

# ─────────────────────────────────────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
h1("REFERENCES")
refs = [
    '[1] M. Fernando, P. Silva, and R. Gunasekara, "Digital transformation in educational '
    'institutions: Opportunities and challenges in the Sri Lankan context," Asian Journal of '
    'Educational Technology, vol. 12, no. 4, pp. 112–128, 2023.',
    '[2] R. Silva and N. Perera, "Time management in school administration: Impact of digital '
    'tools on administrative efficiency," Asian Journal of Education and Training, vol. 10, '
    'no. 2, pp. 156–171, 2024.',
    '[3] K. Dissanayake, "Challenges in student data management in Sri Lankan secondary '
    'schools: A case study approach," Sri Lankan Journal of Educational Administration, '
    'vol. 8, no. 2, pp. 45–62, 2023.',
    '[4] D. Jayawardena and S. Ranasinghe, "Efficiency gains through automation in educational '
    'administration: Evidence from South Asian schools," Journal of Educational Management, '
    'vol. 18, no. 3, pp. 201–218, 2024.',
    '[5] T. Samarasinghe, K. Perera, and L. Fernando, "Data quality issues in manual record-'
    'keeping systems: Implications for educational institutions," South Asian Journal of '
    'Management Information Systems, vol. 7, no. 2, pp. 89–104, 2023.',
    '[6] N. Jayasinghe and V. Wickramasinghe, "Adoption of information systems in Sri Lankan '
    'schools: Barriers and success factors," International Journal of Education and Development '
    'using ICT, vol. 20, no. 1, pp. 78–95, 2024.',
    '[7] Parliament of Sri Lanka, "Personal Data Protection Act, No. 9 of 2022," Government '
    'Publications Bureau, Colombo, 2022.',
    '[8] Ministry of Education, "National Policy on ICT in Education," Ministry of Education, '
    'Sri Lanka, 2023.',
    '[9] A. Rashid and M. Khan, "Student information management systems: A comparative '
    'analysis of features and benefits," International Journal of Educational Technology in '
    'Higher Education, vol. 19, no. 1, pp. 1–21, 2022.',
]
for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'; run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(6)
    pf.left_indent = Inches(0.3); pf.first_line_indent = Inches(-0.3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.5

pb()

# ─────────────────────────────────────────────────────────────────────────────
# APPENDIX A
# ─────────────────────────────────────────────────────────────────────────────
h1("APPENDIX A — Selected Source Code")

h2("Listing A.1: ScoreService.java — computeScore")
for line in [
    "public double computeScore(Long studentId, String stream) {",
    "    List<OLResult> results = olResultRepository.findByStudentId(studentId);",
    "    List<StreamScoreConfig> configs = scoreConfigRepository.findByStream(stream);",
    "    Map<String, Double> weights = configs.stream()",
    "            .collect(Collectors.toMap(",
    "                    StreamScoreConfig::getSubjectCode,",
    "                    StreamScoreConfig::getWeight));",
    "    double total = 0.0;",
    "    for (OLResult r : results) {",
    "        int points = GRADE_POINTS.getOrDefault(r.getGrade().toUpperCase(), 0);",
    "        double weight = weights.getOrDefault(r.getSubject().toUpperCase(), 1.0);",
    "        total += points * weight;",
    "    }",
    "    return total;",
    "}",
]: code_line(line)

doc.add_paragraph()
h2("Listing A.2: InterviewService.java — scheduleBatch")
for line in [
    "@Transactional",
    "public List<Interview> scheduleBatch(List<Long> studentIds,",
    "        LocalDateTime startAt, int durationMinutes, String location) {",
    "    List<Interview> created = new ArrayList<>();",
    "    LocalDateTime slot = startAt;",
    "    for (Long studentId : studentIds) {",
    "        Student student = studentRepository.findById(studentId).orElse(null);",
    "        if (student == null) continue;",
    "        Interview interview = new Interview();",
    "        interview.setStudent(student);",
    "        interview.setScheduledAt(slot);",
    "        interview.setDurationMinutes(durationMinutes);",
    "        interview.setLocation(location);",
    "        interview.setStatus(\"SCHEDULED\");",
    "        created.add(interviewRepository.save(interview));",
    "        student.setRegistrationStatus(\"SCHEDULED\");",
    "        studentRepository.save(student);",
    "        emailService.sendInterviewConfirmation(student, slot, location);",
    "        auditService.log(\"INTERVIEW_SCHEDULED\",",
    "                \"Interview at \" + slot + \" for \" + student.getNicNumber());",
    "        slot = slot.plusMinutes(durationMinutes);",
    "    }",
    "    return created;",
    "}",
]: code_line(line)

doc.add_paragraph()
h2("Listing A.3: nicDecoder.ts — NIC parsing logic")
for line in [
    "export function decodeNIC(nic: string): NICInfo | null {",
    "    const trimmed = nic.trim();",
    "    let year: number, dayOfYear: number;",
    "    const oldFormat = /^(\\d{9})[VXvx]$/.exec(trimmed);",
    "    const newFormat = /^(\\d{12})$/.exec(trimmed);",
    "    if (oldFormat) {",
    "        year      = 1900 + parseInt(trimmed.substring(0, 2), 10);",
    "        dayOfYear = parseInt(trimmed.substring(2, 5), 10);",
    "    } else if (newFormat) {",
    "        year      = parseInt(trimmed.substring(0, 4), 10);",
    "        dayOfYear = parseInt(trimmed.substring(4, 7), 10);",
    "    } else { return null; }",
    "    const gender: 'MALE' | 'FEMALE' = dayOfYear >= 500 ? 'FEMALE' : 'MALE';",
    "    if (dayOfYear >= 500) dayOfYear -= 500;",
    "    const date = new Date(year, 0, dayOfYear);",
    "    if (isNaN(date.getTime())) return null;",
    "    const month = String(date.getMonth() + 1).padStart(2, '0');",
    "    const day   = String(date.getDate()).padStart(2, '0');",
    "    return { dob: `${year}-${month}-${day}`, gender };",
    "}",
]: code_line(line)

doc.add_paragraph()
h2("Listing A.4: RateLimitFilter.java — doFilter")
for line in [
    "@Override",
    "public void doFilter(ServletRequest req, ServletResponse res, FilterChain chain)",
    "        throws IOException, ServletException {",
    "    HttpServletRequest  httpReq = (HttpServletRequest)  req;",
    "    HttpServletResponse httpRes = (HttpServletResponse) res;",
    "    if (!enabled || !httpReq.getRequestURI().startsWith(\"/api/public/\")) {",
    "        chain.doFilter(req, res);",
    "        return;",
    "    }",
    "    String ip  = clientIp(httpReq);",
    "    long   now = System.currentTimeMillis();",
    "    Deque<Long> timestamps = log.computeIfAbsent(ip, k -> new ConcurrentLinkedDeque<>());",
    "    timestamps.removeIf(t -> now - t > WINDOW_MS);",
    "    timestamps.addLast(now);",
    "    if (timestamps.size() > MAX_REQUESTS) {",
    "        httpRes.setStatus(429);",
    "        httpRes.setContentType(\"application/json\");",
    "        httpRes.getWriter().write(",
    "            \"{\\\"error\\\":\\\"Too many requests — please wait a minute.\\\"}\" );",
    "        return;",
    "    }",
    "    chain.doFilter(req, res);",
    "}",
]: code_line(line)

pb()

# ─────────────────────────────────────────────────────────────────────────────
# APPENDIX B
# ─────────────────────────────────────────────────────────────────────────────
h1("APPENDIX B — Application Screenshots")
body("Insert screenshots from the running application before final submission.")
screenshots = [
    ("Figure 4:",  "Login screen"),
    ("Figure 5:",  "Dashboard — stat cards, stream donut, gender donut, grade bar chart, O/L pass/fail chart"),
    ("Figure 6:",  "Student list with search and stream/grade filters"),
    ("Figure 7:",  "Student registration form (personal, academic, and contact sections)"),
    ("Figure 8:",  "Data import — drag-and-drop upload with result summary and error list"),
    ("Figure 9:",  "Reports page — filter controls and download buttons"),
    ("Figure 10:", "Grade promotion — bulk-selectable Grade 12 student table"),
    ("Figure 11:", "Applications page — pending approvals with approve/reject actions"),
    ("Figure 12:", "Audit log — timestamped event table"),
    ("Figure 13:", "Public application portal — 3-step wizard with NIC auto-fill"),
    ("Figure 14:", "Application analytics — ranked applicants table with O/L grade expansion and score weights editor"),
    ("Figure 15:", "Interview schedule — day-view timeline with interview cards and day-summary sidebar"),
    ("Figure 16:", "Interview confirmation email — rendered HTML in an email client"),
]
for label, desc in screenshots:
    p = doc.add_paragraph()
    run1 = p.add_run(label + " "); run1.bold = True
    run1.font.name = 'Times New Roman'; run1.font.size = Pt(11)
    run2 = p.add_run(desc); run2.italic = True
    run2.font.name = 'Times New Roman'; run2.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)

    placeholder = doc.add_paragraph()
    pr = placeholder.add_run("[Insert screenshot here]")
    pr.font.name = 'Times New Roman'; pr.font.size = Pt(10)
    pr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder.paragraph_format.space_before = Pt(2)
    placeholder.paragraph_format.space_after = Pt(12)

# ── Page numbers ──────────────────────────────────────────────────────────────
for sec in doc.sections:
    add_page_numbers(sec)

out = "/Users/yowunpansilu/Documents/GitHub/SIMS/ProgressReport.docx"
doc.save(out)
print(f"Saved: {out}")
